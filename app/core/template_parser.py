"""
Template Parser: Phân tích file .docx để trích xuất placeholders Jinja2.

Giải quyết vấn đề Word split runs:
- Word hay tách {{ ho_ten }} thành nhiều XML run
- Module này merge run trước khi parse → không mất placeholder nào
"""
import re
from pathlib import Path
from typing import Optional
from docxtpl import DocxTemplate
from docx import Document
from app.core.logging import logger

# Regex patterns
VARIABLE_PATTERN = re.compile(r"\{\{\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*\}\}")
FOR_PATTERN = re.compile(r"\{%[-\s]*for\s+(\w+)\s+in\s+(\w+)\s*[-\s]*%\}")
FILTER_PATTERN = re.compile(r"\{\{\s*\w+(?:\.\w+)+\s*\}\}")  # {{ item.field }}
LOOP_FIELD_PATTERN = re.compile(r"\{\{\s*\w+\.([a-zA-Z_][a-zA-Z0-9_]*)\s*\}\}")
# Index-based access: {{ hanh_khach_trai[0].ho_ten }} or {{ hanh_khach[0].ho_ten }}
INDEX_ACCESS_PATTERN = re.compile(r"\{\{\s*([a-zA-Z_][a-zA-Z0-9_]*)\[\d+\]\.([a-zA-Z_][a-zA-Z0-9_]*)\s*\}\}")


def _merge_runs_in_paragraph(paragraph) -> str:
    """
    Ghép tất cả runs trong một paragraph thành 1 string.
    Giải quyết vấn đề Word tách {{ ho_ten }} thành nhiều fragments.
    """
    return "".join(run.text for run in paragraph.runs)


def _extract_all_text(doc: Document) -> list[str]:
    """Trích xuất text từ tất cả paragraphs, bao gồm cả text trong bảng."""
    texts = []

    # Paragraphs thường
    for para in doc.paragraphs:
        texts.append(_merge_runs_in_paragraph(para))

    # Paragraphs trong bảng
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    texts.append(_merge_runs_in_paragraph(para))

    return texts


def parse_template(filepath: str | Path) -> dict:
    """
    Parse template Word và trả về metadata đầy đủ.

    Returns:
        {
            "fields": [{"key": "ho_ten", "type": "text"}],
            "tables": [{"key": "danh_sach", "loop_var": "item", "columns": ["ten", "so_tien"]}],
            "raw_variables": ["ho_ten", "ngay_sinh", ...]
        }
    """
    filepath = Path(filepath)
    if not filepath.exists():
        raise FileNotFoundError(f"Template không tồn tại: {filepath}")

    try:
        doc = Document(str(filepath))
    except Exception as e:
        raise ValueError(f"Không thể đọc file Word: {e}")

    all_texts = _extract_all_text(doc)
    full_text = "\n".join(all_texts)

    # Parse biến đơn (không phải trong loop)
    all_vars = set(VARIABLE_PATTERN.findall(full_text))

    # Parse vòng lặp {% for item in danh_sach %}
    loops = FOR_PATTERN.findall(full_text)

    # Xác định loop vars và loop collections
    loop_vars = {loop_var for loop_var, _ in loops}
    loop_collections = {collection for _, collection in loops}

    # Parse fields con trong loop {{ item.field_name }}
    loop_field_details: dict[str, set] = {}
    for loop_var, collection in loops:
        pattern = re.compile(rf"\{{\{{\s*{loop_var}\.([a-zA-Z_][a-zA-Z0-9_]*)\s*\}}\}}")
        cols = set(pattern.findall(full_text))
        loop_field_details[collection] = cols

    # Detect index-based list access: {{ hanh_khach_trai[0].ho_ten }}
    index_collections: dict[str, set] = {}
    for collection, field in INDEX_ACCESS_PATTERN.findall(full_text):
        index_collections.setdefault(collection, set()).add(field)

    # Merge index-based collections into loop_collections & field details
    for collection, cols in index_collections.items():
        loop_collections.add(collection)
        if collection in loop_field_details:
            loop_field_details[collection].update(cols)
        else:
            loop_field_details[collection] = cols

    # Fields thường = tất cả vars - loop vars - loop collections - index collections
    regular_fields = all_vars - loop_vars - loop_collections

    # Build result
    fields = [
        {"key": key, "type": _infer_field_type(key)}
        for key in sorted(regular_fields)
    ]

    tables = [
        {
            "key": collection,
            "loop_var": next((lv for lv, coll in loops if coll == collection), "item"),
            "columns": sorted(loop_field_details.get(collection, [])),
            "access": "index" if collection in index_collections and collection not in {coll for _, coll in loops} else "loop",
        }
        for collection in sorted(loop_collections)
    ]

    logger.info("template_parsed", filepath=str(filepath), fields=len(fields), tables=len(tables))

    return {
        "fields": fields,
        "tables": tables,
        "raw_variables": sorted(all_vars),
    }


def _infer_field_type(key: str) -> str:
    """Đoán type dựa theo tên field (heuristic)."""
    key_lower = key.lower()

    date_hints = ["ngay", "date", "thang", "nam", "time", "gio", "datetime"]
    number_hints = ["so_", "num", "count", "tong", "gia", "luong", "phi", "tien", "amount", "price", "qty"]
    bool_hints = ["is_", "has_", "co_", "flag", "active", "enable"]

    if any(h in key_lower for h in date_hints):
        return "date"
    if any(h in key_lower for h in number_hints):
        return "number"
    if any(h in key_lower for h in bool_hints):
        return "boolean"
    return "text"
